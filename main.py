"""
BMC RAS测试报告深度分析系统 - PyQt6桌面版
支持三轮AI推理 + 跨测试关联分析
"""

import sys
import os
import json
import re
import tempfile
import shutil
import traceback
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QLineEdit, QTextEdit, QFileDialog,
    QProgressBar, QGroupBox, QSpinBox, QMessageBox, QTabWidget,
    QSplitter, QComboBox, QStatusBar
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt6.QtGui import QFont, QTextCursor, QIcon

import docx
from paddleocr import PaddleOCR
import requests


class ComponentType(Enum):
    """部件类型"""
    CPU = "CPU"
    MEMORY = "Memory"
    PCIE = "PCIe"
    SATA = "SATA"
    OTHER = "Other"


class ErrorType(Enum):
    """错误类型"""
    CORRECTABLE = "Correctable"
    UNCORRECTABLE = "Uncorrectable"
    FATAL = "Fatal"
    OTHER = "Other"


@dataclass
class TestCase:
    """测试用例数据结构"""
    title: str
    component: ComponentType
    error_type: ErrorType
    test_result: str
    raw_text: str
    ocr_images: List[Dict]
    register_values: Dict[str, str] = None
    silk_screen: str = None

    def __post_init__(self):
        if self.register_values is None:
            self.register_values = {}
        if self.ocr_images is None:
            self.ocr_images = []


class BMCRASAnalyzer:
    """BMC RAS分析器核心"""

    def __init__(self, model_name: str = "qwen3:8b", ollama_url: str = "http://localhost:11434/api/generate"):
        self.model_name = model_name
        self.ollama_url = ollama_url

        # 初始化OCR
        try:
            self.ocr = PaddleOCR(use_angle_cls=True, lang='ch')
        except:
            self.ocr = PaddleOCR(lang='ch')

    def extract_images_from_docx(self, doc_path: str, output_dir: str) -> List[Tuple[str, int]]:
        """从Word文档提取图片"""
        doc = docx.Document(doc_path)
        image_list = []
        os.makedirs(output_dir, exist_ok=True)

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                img_name = os.path.basename(rel.target_ref)
                img_path = os.path.join(output_dir, img_name)

                with open(img_path, 'wb') as f:
                    f.write(img_data)

                para_index = self._find_image_paragraph(doc, rel.rId)
                image_list.append((img_path, para_index))

        return sorted(image_list, key=lambda x: x[1])

    def _find_image_paragraph(self, doc, rel_id: str) -> int:
        for idx, para in enumerate(doc.paragraphs):
            if rel_id in para._element.xml:
                return idx
        return -1

    def ocr_images(self, image_paths: List[str]) -> Dict[str, Dict]:
        """OCR识别"""
        results = {}
        for img_path in image_paths:
            try:
                result = self.ocr.ocr(img_path, cls=True)
                text_lines = []

                if result and result[0]:
                    for line in result[0]:
                        if line and len(line) >= 2:
                            text_lines.append(line[1][0])

                full_text = "\n".join(text_lines)
                results[img_path] = {
                    'text': full_text,
                    'lines': text_lines,
                    'registers': self._extract_registers(full_text),
                    'silk_screen': self._extract_silk_screen(full_text)
                }
            except Exception as e:
                results[img_path] = {'text': f'OCR失败: {str(e)}', 'lines': [], 'registers': {}, 'silk_screen': None}

        return results

    def _extract_registers(self, text: str) -> Dict[str, str]:
        """提取寄存器"""
        registers = {}
        patterns = [
            r'(MCA_BANK\d+|MCi_\w+)\s*[=:]\s*(0x[0-9a-fA-F]+)',
            r'(DIMM[A-Z]?\d+)\s*[=:]\s*([^\s]+)',
            r'([A-Z_]+_REG(?:ISTER)?)\s*[=:]\s*(0x[0-9a-fA-F]+|\d+)',
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                registers[match[0]] = match[1]
        return registers

    def _extract_silk_screen(self, text: str) -> Optional[str]:
        """提取丝印"""
        patterns = [
            r'(silk.*?screen|丝印)[:\s]+([A-Z0-9_-]+)',
            r'(label|标签)[:\s]+([A-Z0-9_-]+)',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(2)
        return None

    def parse_test_structure(self, doc_path: str) -> List[TestCase]:
        """解析测试结构"""
        doc = docx.Document(doc_path)
        test_cases = []
        current_case = None
        current_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_title = self._is_test_title(text, para)

            if is_title:
                if current_case:
                    current_case.raw_text = "\n".join(current_text)
                    test_cases.append(current_case)

                current_case = TestCase(
                    title=text,
                    component=self._identify_component(text),
                    error_type=self._identify_error_type(text),
                    test_result=self._extract_result(text),
                    raw_text="",
                    ocr_images=[]
                )
                current_text = [text]
            elif current_case:
                current_text.append(text)

        if current_case:
            current_case.raw_text = "\n".join(current_text)
            test_cases.append(current_case)

        return test_cases

    def _is_test_title(self, text: str, para) -> bool:
        is_heading = para.style.name.startswith('Heading')
        is_bold = para.runs and para.runs[0].bold if para.runs else False
        has_keywords = any(k in text.lower() for k in ['test', '测试', 'cpu', 'memory', 'pcie', 'sata', 'inject'])
        return (is_heading or (is_bold and len(text) < 150)) and has_keywords

    def _identify_component(self, text: str) -> ComponentType:
        t = text.lower()
        if 'cpu' in t or 'processor' in t:
            return ComponentType.CPU
        elif 'mem' in t or 'dimm' in t:
            return ComponentType.MEMORY
        elif 'pcie' in t:
            return ComponentType.PCIE
        elif 'sata' in t or 'disk' in t:
            return ComponentType.SATA
        return ComponentType.OTHER

    def _identify_error_type(self, text: str) -> ErrorType:
        t = text.lower()
        if 'ce' in t or 'correct' in t:
            return ErrorType.CORRECTABLE
        elif 'uce' in t or 'uncorrect' in t:
            return ErrorType.UNCORRECTABLE
        elif 'fatal' in t:
            return ErrorType.FATAL
        return ErrorType.OTHER

    def _extract_result(self, text: str) -> str:
        if 'pass' in text.lower():
            return "PASS"
        elif 'fail' in text.lower():
            return "FAIL"
        return "UNKNOWN"

    def integrate_ocr(self, test_cases: List[TestCase], image_list, ocr_results) -> List[TestCase]:
        """整合OCR结果"""
        for img_path, para_idx in image_list:
            ocr_data = ocr_results.get(img_path, {})
            for tc in test_cases:
                tc.ocr_images.append({'path': img_path, 'para_idx': para_idx, 'ocr_data': ocr_data})
                if ocr_data.get('registers'):
                    tc.register_values.update(ocr_data['registers'])
                if ocr_data.get('silk_screen') and not tc.silk_screen:
                    tc.silk_screen = ocr_data['silk_screen']
        return test_cases

    def call_ollama(self, prompt: str, timeout: int = 300) -> str:
        """调用Ollama"""
        response = requests.post(
            self.ollama_url,
            json={"model": self.model_name, "prompt": prompt, "stream": False},
            timeout=timeout
        )
        if response.status_code == 200:
            return response.json().get('response', '生成失败')
        raise Exception(f"API错误: {response.status_code}")

    def extract_structured_info(self, tc: TestCase) -> Dict:
        """第一轮：提取结构化信息"""
        prompt = f"""从BMC RAS测试数据中提取关键信息。

测试: {tc.title}
部件: {tc.component.value}
错误类型: {tc.error_type.value}

文本:
{tc.raw_text[:600]}

OCR内容:
"""
        for idx, img in enumerate(tc.ocr_images[:2]):
            prompt += f"\n图{idx + 1}: {img['ocr_data'].get('text', '')[:400]}\n"

        prompt += """
以JSON格式输出:
{"registers": {}, "silk_screen": "", "error_log": "", "diagnosis": "", "position": ""}
"""
        try:
            resp = self.call_ollama(prompt, 120)
            match = re.search(r'\{.*\}', resp, re.DOTALL)
            return json.loads(match.group()) if match else {'raw': resp}
        except:
            return {'error': '提取失败'}

    def analyze_testcase(self, tc: TestCase, info: Dict) -> Dict:
        """第二轮：分析单个测试"""
        prompt = f"""审查BMC RAS测试用例。

测试: {tc.title}
部件: {tc.component.value}
结果: {tc.test_result}

提取信息:
{json.dumps(info, indent=2, ensure_ascii=False)}

原文:
{tc.raw_text[:500]}

检查要点:
1. 寄存器置位是否正确
2. 丝印信息是否完整
3. 诊断结果是否准确
4. 错误日志是否完整
5. 部件位置是否准确

输出格式:
**问题**: (如有)
- [高/中/低] 具体问题
**证据**: 数据支持
**建议**: 改进建议

无问题则输出: "未发现明显问题"
"""
        try:
            resp = self.call_ollama(prompt, 180)
            return {'title': tc.title, 'analysis': resp, 'has_issues': '未发现明显问题' not in resp}
        except Exception as e:
            return {'title': tc.title, 'analysis': f'分析失败: {e}', 'has_issues': False}

    def cross_analysis(self, test_cases: List[TestCase], analyses: List[Dict]) -> Dict:
        """第三轮：关联分析"""
        by_comp = {}
        for tc in test_cases:
            comp = tc.component.value
            by_comp.setdefault(comp, []).append(tc)

        prompt = f"""跨测试用例关联分析。

测试总数: {len(test_cases)}
部件分布: {', '.join(f'{k}: {len(v)}个' for k, v in by_comp.items())}

各部件测试:
"""
        for comp, cases in list(by_comp.items())[:5]:
            prompt += f"\n{comp}:\n"
            for tc in cases[:3]:
                prompt += f"  - {tc.title[:40]} [{tc.test_result}]\n"

        prompt += "\n单项分析摘要:\n"
        for a in analyses[:5]:
            prompt += f"\n{a['title'][:30]}: {a['analysis'][:120]}...\n"

        prompt += """
关联分析要点:
1. 同部件不同错误类型对比
2. 相似测试一致性
3. 丝印连续性
4. 测试覆盖度

输出:
**关联问题**: 跨测试问题
**一致性**: 对比结果
**风险**: 高/中/低
**建议**: 改进建议
"""
        try:
            resp = self.call_ollama(prompt, 300)
            return {'analysis': resp, 'summary': by_comp}
        except Exception as e:
            return {'analysis': f'失败: {e}', 'summary': by_comp}

    def generate_report(self, test_cases, analyses, cross) -> str:
        """生成报告"""
        total = len(test_cases)
        passed = sum(1 for tc in test_cases if tc.test_result == "PASS")
        issues = sum(1 for a in analyses if a.get('has_issues', False))

        report = f"""# 🔬 BMC RAS测试报告 - 深度审查分析

## 📊 执行摘要

- 测试用例总数: {total}
- PASS数量: {passed} ({passed / total * 100:.1f}%)
- 发现问题用例: {issues}
- 问题发现率: {issues / total * 100:.1f}%

### 部件测试覆盖度

"""
        comp_stats = {}
        for tc in test_cases:
            comp = tc.component.value
            comp_stats[comp] = comp_stats.get(comp, 0) + 1

        for comp, count in sorted(comp_stats.items()):
            report += f"- **{comp}**: {count} 个测试\n"

        report += "\n---\n\n## 🔍 单项测试详细分析\n\n"

        for idx, a in enumerate(analyses, 1):
            status = "⚠️ **发现问题**" if a.get('has_issues') else "✅ **通过审查**"
            report += f"### {idx}. {a['title']}\n\n{status}\n\n{a['analysis']}\n\n---\n\n"

        report += f"## 🔗 跨测试用例关联分析\n\n{cross.get('analysis', '未执行')}\n\n---\n\n"

        report += "## 📋 总体结论\n\n"
        if issues > total * 0.3:
            report += "⚠️ **风险等级**: 高\n\n"
        elif issues > total * 0.1:
            report += "⚡ **风险等级**: 中\n\n"
        else:
            report += "✅ **风险等级**: 低\n\n"

        report += """### 建议措施

1. 复测所有高严重度问题用例
2. 补充缺失的测试场景
3. 优化测试流程和文档规范
"""
        return report


class AnalyzerThread(QThread):
    """分析线程"""
    progress_signal = pyqtSignal(int, str)
    result_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)

    def __init__(self, doc_path, model, max_cases):
        super().__init__()
        self.doc_path = doc_path
        self.model = model
        self.max_cases = max_cases

    def run(self):
        try:
            temp_dir = tempfile.mkdtemp()
            img_dir = os.path.join(temp_dir, "images")

            self.progress_signal.emit(0, "🚀 初始化分析器...")
            analyzer = BMCRASAnalyzer(model_name=self.model)

            self.progress_signal.emit(10, "📷 提取图片...")
            image_list = analyzer.extract_images_from_docx(self.doc_path, img_dir)

            if not image_list:
                self.error_signal.emit("未找到图片")
                return

            self.progress_signal.emit(20, "🔍 OCR识别...")
            img_paths = [img[0] for img in image_list]
            ocr_results = analyzer.ocr_images(img_paths)

            self.progress_signal.emit(40, "📋 解析测试用例...")
            test_cases = analyzer.parse_test_structure(self.doc_path)

            if not test_cases:
                self.error_signal.emit("未识别到测试用例")
                return

            self.progress_signal.emit(50, "🔗 整合数据...")
            test_cases = analyzer.integrate_ocr(test_cases, image_list, ocr_results)

            # 限制数量
            cases_to_analyze = test_cases[:self.max_cases]

            self.progress_signal.emit(60, "🧠 AI提取信息...")
            infos = []
            for idx, tc in enumerate(cases_to_analyze):
                self.progress_signal.emit(60 + int(10 * idx / len(cases_to_analyze)),
                                          f"提取 {idx + 1}/{len(cases_to_analyze)}")
                infos.append(analyzer.extract_structured_info(tc))

            self.progress_signal.emit(70, "🔬 AI分析测试...")
            analyses = []
            for idx, (tc, info) in enumerate(zip(cases_to_analyze, infos)):
                self.progress_signal.emit(70 + int(15 * idx / len(cases_to_analyze)),
                                          f"分析 {idx + 1}/{len(cases_to_analyze)}")
                analyses.append(analyzer.analyze_testcase(tc, info))

            self.progress_signal.emit(90, "🔗 关联分析...")
            cross = analyzer.cross_analysis(cases_to_analyze, analyses)

            self.progress_signal.emit(95, "📄 生成报告...")
            report = analyzer.generate_report(cases_to_analyze, analyses, cross)

            shutil.rmtree(temp_dir)
            self.progress_signal.emit(100, "✅ 完成!")
            self.result_signal.emit(report)

        except Exception as e:
            self.error_signal.emit(f"处理出错: {str(e)}\n\n{traceback.format_exc()}")


class MainWindow(QMainWindow):
    """主窗口"""

    def __init__(self):
        super().__init__()
        self.init_ui()
        self.analyzer_thread = None
        self.doc_path = None

    def init_ui(self):
        self.setWindowTitle("🔬 BMC RAS测试报告深度分析系统")
        self.setGeometry(100, 100, 1400, 900)

        # 主widget
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QHBoxLayout(main_widget)

        # 左侧控制面板
        left_panel = self.create_control_panel()

        # 右侧结果区域
        right_panel = self.create_result_panel()

        # 分割器
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)

        main_layout.addWidget(splitter)

        # 状态栏
        self.statusBar().showMessage("就绪")

        # 样式
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QGroupBox {
                font-weight: bold;
                border: 2px solid #cccccc;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 10px;
                border-radius: 5px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:pressed {
                background-color: #3d8b40;
            }
            QPushButton:disabled {
                background-color: #cccccc;
                color: #666666;
            }
            QLineEdit {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
            }
            QLineEdit:focus {
                border: 2px solid #4CAF50;
            }
        """)

    def create_control_panel(self):
        """创建控制面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)

        # 标题
        title = QLabel("🔬 BMC RAS分析系统")
        title.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)

        # 文件选择组
        file_group = QGroupBox("📄 测试报告")
        file_layout = QVBoxLayout()

        self.file_label = QLabel("未选择文件")
        self.file_label.setWordWrap(True)
        self.file_label.setStyleSheet("padding: 5px; background: white; border-radius: 3px;")
        file_layout.addWidget(self.file_label)

        select_btn = QPushButton("📁 选择Word文档")
        select_btn.clicked.connect(self.select_file)
        file_layout.addWidget(select_btn)

        file_group.setLayout(file_layout)
        layout.addWidget(file_group)

        # 模型配置组
        model_group = QGroupBox("🤖 模型配置")
        model_layout = QVBoxLayout()

        model_layout.addWidget(QLabel("Ollama模型:"))
        self.model_combo = QComboBox()
        self.model_combo.addItems(["qwen3:8b", "qwen2.5:7b", "deepseek-r1:7b", "llama3.1:8b"])
        self.model_combo.setEditable(True)
        model_layout.addWidget(self.model_combo)

        model_layout.addWidget(QLabel("最大分析用例数:"))
        self.max_cases_spin = QSpinBox()
        self.max_cases_spin.setRange(5, 50)
        self.max_cases_spin.setValue(10)
        self.max_cases_spin.setSuffix(" 个")
        model_layout.addWidget(self.max_cases_spin)

        model_group.setLayout(model_layout)
        layout.addWidget(model_group)

        # 进度组
        progress_group = QGroupBox("📊 分析进度")
        progress_layout = QVBoxLayout()

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        progress_layout.addWidget(self.progress_bar)

        self.progress_label = QLabel("等待开始...")
        self.progress_label.setWordWrap(True)
        progress_layout.addWidget(self.progress_label)

        progress_group.setLayout(progress_layout)
        layout.addWidget(progress_group)

        # 开始按钮
        self.start_btn = QPushButton("🚀 开始深度分析")
        self.start_btn.setMinimumHeight(50)
        self.start_btn.clicked.connect(self.start_analysis)
        self.start_btn.setEnabled(False)
        layout.addWidget(self.start_btn)

        # 说明
        info_group = QGroupBox("ℹ️ 系统说明")
        info_layout = QVBoxLayout()

        info_text = QTextEdit()
        info_text.setReadOnly(True)
        info_text.setMaximumHeight(250)
        info_text.setHtml("""
<b>分析流程:</b><br>
1. <b>OCR识别</b> → 提取寄存器、丝印、日志<br>
2. <b>结构化提取</b> → AI第一轮提取<br>
3. <b>单项分析</b> → AI第二轮审查<br>
4. <b>关联分析</b> → AI第三轮对比<br>
5. <b>综合报告</b> → 生成问题清单<br><br>

<b>审查重点:</b><br>
✓ 寄存器置位正确性<br>
✓ 丝印信息完整性<br>
✓ 诊断结果准确性<br>
✓ 跨测试一致性<br>
✓ 部件测试覆盖度<br><br>

<b>系统要求:</b><br>
• Ollama运行在本地<br>
• 模型已下载<br>
• Word文档格式正确<br>
        """)
        info_layout.addWidget(info_text)

        info_group.setLayout(info_layout)
        layout.addWidget(info_group)

        layout.addStretch()
        return panel

    def create_result_panel(self):
        """创建结果面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)

        # 标签栏
        self.tab_widget = QTabWidget()

        # 分析报告标签
        self.report_text = QTextEdit()
        self.report_text.setReadOnly(True)
        self.report_text.setPlaceholderText("分析报告将在此显示...")

        # 设置等宽字体
        font = QFont("Consolas", 10)
        self.report_text.setFont(font)

        self.tab_widget.addTab(self.report_text, "📄 分析报告")

        # 日志标签
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(200)
        log_font = QFont("Courier", 9)
        self.log_text.setFont(log_font)

        self.tab_widget.addTab(self.log_text, "📝 运行日志")

        layout.addWidget(self.tab_widget)

        # 操作按钮
        button_layout = QHBoxLayout()

        self.save_btn = QPushButton("💾 保存报告")
        self.save_btn.clicked.connect(self.save_report)
        self.save_btn.setEnabled(False)
        button_layout.addWidget(self.save_btn)

        self.copy_btn = QPushButton("📋 复制报告")
        self.copy_btn.clicked.connect(self.copy_report)
        self.copy_btn.setEnabled(False)
        button_layout.addWidget(self.copy_btn)

        self.clear_btn = QPushButton("🗑️ 清空")
        self.clear_btn.clicked.connect(self.clear_results)
        button_layout.addWidget(self.clear_btn)

        layout.addLayout(button_layout)

        return panel

    def select_file(self):
        """选择文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择测试报告",
            "",
            "Word文档 (*.docx)"
        )

        if file_path:
            self.doc_path = file_path
            self.file_label.setText(f"📄 {Path(file_path).name}")
            self.start_btn.setEnabled(True)
            self.log(f"✓ 已选择文件: {file_path}")

    def start_analysis(self):
        """开始分析"""
        if not self.doc_path:
            QMessageBox.warning(self, "警告", "请先选择测试报告文件!")
            return

        model = self.model_combo.currentText()
        max_cases = self.max_cases_spin.value()

        # 禁用按钮
        self.start_btn.setEnabled(False)
        self.save_btn.setEnabled(False)
        self.copy_btn.setEnabled(False)

        # 清空结果
        self.report_text.clear()
        self.progress_bar.setValue(0)

        self.log(f"开始分析，模型: {model}, 最大用例数: {max_cases}")

        # 创建并启动线程
        self.analyzer_thread = AnalyzerThread(self.doc_path, model, max_cases)
        self.analyzer_thread.progress_signal.connect(self.update_progress)
        self.analyzer_thread.result_signal.connect(self.show_result)
        self.analyzer_thread.error_signal.connect(self.show_error)
        self.analyzer_thread.start()

    def update_progress(self, value, message):
        """更新进度"""
        self.progress_bar.setValue(value)
        self.progress_label.setText(message)
        self.log(f"[{value}%] {message}")

    def show_result(self, report):
        """显示结果"""
        self.report_text.setPlainText(report)
        self.save_btn.setEnabled(True)
        self.copy_btn.setEnabled(True)
        self.start_btn.setEnabled(True)
        self.log("✓ 分析完成!")
        self.statusBar().showMessage("分析完成", 5000)
        QMessageBox.information(self, "完成", "分析已完成，请查看报告!")

    def show_error(self, error):
        """显示错误"""
        self.report_text.setPlainText(f"❌ 错误:\n\n{error}")
        self.start_btn.setEnabled(True)
        self.log(f"✗ 错误: {error}")
        QMessageBox.critical(self, "错误", f"分析失败:\n{error[:200]}...")

    def save_report(self):
        """保存报告"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存分析报告",
            "BMC_RAS_分析报告.md",
            "Markdown (*.md);;文本文件 (*.txt)"
        )

        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.report_text.toPlainText())
                self.log(f"✓ 报告已保存: {file_path}")
                QMessageBox.information(self, "成功", "报告已保存!")
            except Exception as e:
                QMessageBox.critical(self, "错误", f"保存失败: {e}")

    def copy_report(self):
        """复制报告"""
        clipboard = QApplication.clipboard()
        clipboard.setText(self.report_text.toPlainText())
        self.log("✓ 报告已复制到剪贴板")
        self.statusBar().showMessage("已复制到剪贴板", 3000)

    def clear_results(self):
        """清空结果"""
        self.report_text.clear()
        self.log_text.clear()
        self.progress_bar.setValue(0)
        self.progress_label.setText("等待开始...")
        self.log("已清空结果")

    def log(self, message):
        """添加日志"""
        from datetime import datetime
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.append(f"[{timestamp}] {message}")
        self.log_text.moveCursor(QTextCursor.MoveOperation.End)


def main():
    app = QApplication(sys.argv)
    app.setApplicationName("BMC RAS分析系统")

    window = MainWindow()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()
